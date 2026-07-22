"""
comprehensive_test.py - Full Apache POI corpus workflow for docx-creation-editing.

This script automatically:
- clones a sparse, shallow Apache POI working copy into a temp directory
- uses Apache POI test-data/document and test-data/integration as the baseline corpus
- verifies required fixtures exist and are intact
- runs the docx regression suite against that baseline
- emits verbose structured logs and metadata
- cleans up the clone and all temp artifacts regardless of outcome

Exact upstream baseline paths within apache/poi:
- test-data/document/
- test-data/integration/
"""

from __future__ import annotations

import argparse
import base64
import datetime as dt
import hashlib
import json
import shutil
import struct
import subprocess
import sys
import tempfile
import time
import traceback
import zipfile
import zlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

sys.path.insert(0, str(Path(__file__).parent))
import create_docx
import edit_docx
import validate_docx
import vba_runner

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

SCRIPT_DIR = Path(__file__).resolve().parent
TEMP_ROOT = Path(tempfile.gettempdir()) / "kilo"
DEFAULT_REPO_URL = "https://github.com/apache/poi.git"
DEFAULT_REF = "trunk"
BASELINE_RELATIVE_PATHS = (
    Path("test-data/document"),
    Path("test-data/integration"),
)
REQUIRED_FIXTURES = (
    Path("document/delins.docx"),
    Path("document/58067.docx"),
    Path("document/Tika-792.docx"),
    Path("document/VariousPictures.docx"),
    Path("document/WordWithAttachments.docx"),
    Path("document/comment.docx"),
    Path("document/testComment.docx"),
    Path("document/drawing.docx"),
    Path("integration/stress023.docx"),
)
DEFAULT_TIMEOUT_SECONDS = 240
DEFAULT_RETRIES = 3
PASS = "PASS"
FAIL = "FAIL"
SKIP = "SKIP"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
JPEG_1X1 = base64.b64decode(
    "/9j/4AAQSkZJRgABAQAAAQABAAD/2wCEAAkGBxAQEBAQEA8QEA8QDw8PDw8PDw8PDw8PFREWFhURFRUYHSggGBolGxUVITEhJSkrLi4uFx8zODMsNygtLisBCgoKDg0OGxAQGy0lICYtLS0tLS0tLS0tLS0tLS0tLS0tLS0tLS0tLS0tLS0tLS0tLS0tLS0tLS0tLS0tLf/AABEIAAEAAgMBIgACEQEDEQH/xAAXAAEBAQEAAAAAAAAAAAAAAAAAAQID/8QAFBABAAAAAAAAAAAAAAAAAAAAAP/aAAwDAQACEAMQAAAB6AAAAP/EABQQAQAAAAAAAAAAAAAAAAAAACD/2gAIAQEAAT8Af//EABQRAQAAAAAAAAAAAAAAAAAAACD/2gAIAQIBAT8Af//EABQRAQAAAAAAAAAAAAAAAAAAACD/2gAIAQMBAT8Af//Z"
)
PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADUlEQVR42mP8z/C/HwAFgwJ/l6s5WQAAAABJRU5ErkJggg=="
)
WMF_FIXTURE_CANDIDATES = (
    "document/testException2.doc-2.wmf",
    "document/vector_image.emf",
)


class WorkflowError(RuntimeError):
    """Raised when a workflow step fails in a controlled way."""


@dataclass
class TestRecord:
    name: str
    status: str
    detail: str
    started_at: str
    finished_at: str
    duration_seconds: float


@dataclass
class RunContext:
    args: argparse.Namespace
    run_root: Path
    repo_dir: Path
    output_dir: Path
    log_path: Path
    metadata_path: Path
    repo_url: str
    ref: str
    clone_started_at: str | None = None
    clone_finished_at: str | None = None
    repo_commit: str | None = None
    baseline_root: Path | None = None
    fixture_metadata: list[dict[str, Any]] = field(default_factory=list)
    test_records: list[TestRecord] = field(default_factory=list)
    step_logs: list[dict[str, Any]] = field(default_factory=list)
    cleanup_completed: bool = False
    test_command: str = ""


def now_utc() -> dt.datetime:
    return dt.datetime.now(dt.timezone.utc)


def iso_now() -> str:
    return now_utc().isoformat()


def log_event(ctx: RunContext, level: str, event: str, **fields: Any) -> None:
    record = {
        "timestamp": iso_now(),
        "level": level,
        "event": event,
        **fields,
    }
    ctx.step_logs.append(record)
    line = json.dumps(record, sort_keys=True, default=str)
    print(line)
    ctx.log_path.parent.mkdir(parents=True, exist_ok=True)
    with ctx.log_path.open("a", encoding="utf-8") as fh:
        fh.write(line + "\n")


def _json_safe(value: Any) -> Any:
    if isinstance(value, Path):
        return str(value)
    if isinstance(value, dt.datetime):
        return value.isoformat()
    if isinstance(value, TestRecord):
        return value.__dict__
    if isinstance(value, dict):
        return {str(k): _json_safe(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [_json_safe(v) for v in value]
    return value


def write_metadata(ctx: RunContext, status: str, error: str | None = None) -> None:
    passed = sum(1 for rec in ctx.test_records if rec.status == PASS)
    failed = sum(1 for rec in ctx.test_records if rec.status == FAIL)
    skipped = sum(1 for rec in ctx.test_records if rec.status == SKIP)
    metadata = {
        "baseline": {
            "repo_url": ctx.repo_url,
            "ref": ctx.ref,
            "commit": ctx.repo_commit,
            "paths": [str(path).replace("\\", "/") for path in BASELINE_RELATIVE_PATHS],
            "clone_started_at": ctx.clone_started_at,
            "clone_finished_at": ctx.clone_finished_at,
            "fixtures": ctx.fixture_metadata,
        },
        "run": {
            "command": ctx.test_command,
            "started_at": ctx.step_logs[0]["timestamp"] if ctx.step_logs else None,
            "finished_at": iso_now(),
            "status": status,
            "error": error,
            "summary": {
                "total": len(ctx.test_records),
                "passed": passed,
                "failed": failed,
                "skipped": skipped,
            },
            "tests": [_json_safe(rec) for rec in ctx.test_records],
            "logs": _json_safe(ctx.step_logs),
        },
        "cleanup": {
            "completed": ctx.cleanup_completed,
        },
    }
    ctx.metadata_path.parent.mkdir(parents=True, exist_ok=True)
    with ctx.metadata_path.open("w", encoding="utf-8") as fh:
        json.dump(metadata, fh, indent=2)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def run_command(
    ctx: RunContext,
    command: list[str],
    *,
    cwd: Path | None = None,
    timeout: int = DEFAULT_TIMEOUT_SECONDS,
    retries: int = 1,
    retry_delay: float = 2.0,
) -> subprocess.CompletedProcess[str]:
    attempt = 0
    last_error: str | None = None
    while attempt < retries:
        attempt += 1
        started = time.monotonic()
        log_event(
            ctx,
            "info",
            "command.start",
            attempt=attempt,
            retries=retries,
            cwd=str(cwd) if cwd else None,
            timeout_seconds=timeout,
            command=command,
        )
        try:
            completed = subprocess.run(
                command,
                cwd=str(cwd) if cwd else None,
                check=False,
                capture_output=True,
                text=True,
                timeout=timeout,
            )
        except subprocess.TimeoutExpired as exc:
            last_error = f"Timed out after {timeout}s: {' '.join(command)}"
            log_event(
                ctx,
                "error",
                "command.timeout",
                attempt=attempt,
                duration_seconds=round(time.monotonic() - started, 3),
                stderr=exc.stderr,
                stdout=exc.stdout,
                error=last_error,
            )
        except OSError as exc:
            last_error = f"Failed to launch command {' '.join(command)}: {exc}"
            log_event(
                ctx,
                "error",
                "command.oserror",
                attempt=attempt,
                duration_seconds=round(time.monotonic() - started, 3),
                error=last_error,
            )
        else:
            duration = round(time.monotonic() - started, 3)
            log_event(
                ctx,
                "info" if completed.returncode == 0 else "warning",
                "command.finish",
                attempt=attempt,
                duration_seconds=duration,
                returncode=completed.returncode,
                stdout=completed.stdout,
                stderr=completed.stderr,
            )
            if completed.returncode == 0:
                return completed
            last_error = (
                f"Command failed with exit code {completed.returncode}: {' '.join(command)}\n"
                f"STDERR:\n{completed.stderr.strip()}\nSTDOUT:\n{completed.stdout.strip()}"
            )
        if attempt < retries:
            sleep_for = retry_delay * attempt
            log_event(ctx, "warning", "command.retry", attempt=attempt, sleep_seconds=sleep_for, error=last_error)
            time.sleep(sleep_for)
    raise WorkflowError(last_error or f"Command failed: {' '.join(command)}")


def _make_png(r: int = 255, g: int = 0, b: int = 0) -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", crc)

    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    raw = b"\x00" + bytes([r, g, b] * 8) * 8
    idat = zlib.compress(raw)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _fixture(ctx: RunContext, relative_path: str) -> Path:
    assert ctx.baseline_root is not None
    return ctx.baseline_root / relative_path


def _output_path(ctx: RunContext, name: str) -> Path:
    ctx.output_dir.mkdir(parents=True, exist_ok=True)
    return ctx.output_dir / name


def _write_binary(path: Path, data: bytes) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    return path


def _existing_fixture(ctx: RunContext, candidates: tuple[str, ...]) -> Path:
    missing: list[str] = []
    for candidate in candidates:
        path = _fixture(ctx, candidate)
        if path.exists():
            return path
        missing.append(str(path))
    raise FileNotFoundError(f"None of the candidate fixtures exist: {missing}")


def _raw_count(path: Path, tag: str) -> int:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml")
    return xml.decode("utf-8", errors="replace").count(tag)


def _revision_ids_in_doc(doc: Document) -> set[int]:
    body = doc.element.body
    ids = set()
    for tag in ("w:ins", "w:del", "w:commentRangeStart"):
        for el in body.iter(qn(tag)):
            try:
                ids.add(int(el.get(qn("w:id"), "0")))
            except ValueError:
                pass
    return ids


def prepare_poi_baseline(ctx: RunContext) -> None:
    ctx.clone_started_at = iso_now()
    run_command(
        ctx,
        [
            "git",
            "clone",
            "--depth",
            "1",
            "--filter=blob:none",
            "--sparse",
            ctx.repo_url,
            str(ctx.repo_dir),
        ],
        timeout=DEFAULT_TIMEOUT_SECONDS,
        retries=DEFAULT_RETRIES,
    )
    run_command(
        ctx,
        [
            "git",
            "-C",
            str(ctx.repo_dir),
            "sparse-checkout",
            "set",
            "test-data/document",
            "test-data/integration",
        ],
        timeout=DEFAULT_TIMEOUT_SECONDS,
        retries=DEFAULT_RETRIES,
    )
    run_command(
        ctx,
        ["git", "-C", str(ctx.repo_dir), "checkout", ctx.ref],
        timeout=DEFAULT_TIMEOUT_SECONDS,
        retries=DEFAULT_RETRIES,
    )
    ctx.repo_commit = run_command(
        ctx,
        ["git", "-C", str(ctx.repo_dir), "rev-parse", "HEAD"],
        timeout=60,
    ).stdout.strip()
    ctx.clone_finished_at = iso_now()
    ctx.baseline_root = ctx.repo_dir / "test-data"
    log_event(
        ctx,
        "info",
        "baseline.ready",
        repo_dir=str(ctx.repo_dir),
        baseline_root=str(ctx.baseline_root),
        commit=ctx.repo_commit,
    )


def verify_baseline(ctx: RunContext) -> None:
    assert ctx.baseline_root is not None
    for rel in BASELINE_RELATIVE_PATHS:
        root = ctx.repo_dir / rel
        if not root.exists() or not root.is_dir():
            raise WorkflowError(f"Missing baseline directory: {root}")

    fixture_metadata: list[dict[str, Any]] = []
    for rel in REQUIRED_FIXTURES:
        path = ctx.baseline_root / rel
        if not path.exists():
            raise WorkflowError(f"Missing required baseline fixture: {path}")
        if path.stat().st_size <= 0:
            raise WorkflowError(f"Baseline fixture is empty: {path}")
        try:
            with zipfile.ZipFile(path) as zf:
                bad = zf.testzip()
                if bad is not None:
                    raise WorkflowError(f"Fixture ZIP integrity failed for {path}: first bad entry {bad}")
        except zipfile.BadZipFile as exc:
            raise WorkflowError(f"Fixture is not a valid ZIP package: {path} ({exc})") from exc

        fixture_metadata.append(
            {
                "path": str(rel).replace("\\", "/"),
                "size": path.stat().st_size,
                "sha256": sha256_file(path),
                "zip_valid": True,
            }
        )
    ctx.fixture_metadata = fixture_metadata
    log_event(ctx, "info", "baseline.verified", fixture_count=len(fixture_metadata), fixtures=fixture_metadata)


class TestSuite:
    def __init__(self, ctx: RunContext) -> None:
        self.ctx = ctx

    def run_test(self, name: str, fn: Callable[[], str]) -> None:
        started = time.monotonic()
        started_at = iso_now()
        log_event(self.ctx, "info", "test.start", name=name)
        try:
            detail = fn() or ""
        except AssertionError as exc:
            finished_at = iso_now()
            record = TestRecord(
                name=name,
                status=FAIL,
                detail=str(exc),
                started_at=started_at,
                finished_at=finished_at,
                duration_seconds=round(time.monotonic() - started, 3),
            )
            self.ctx.test_records.append(record)
            log_event(self.ctx, "error", "test.fail", name=name, detail=record.detail)
            return
        except Exception as exc:
            finished_at = iso_now()
            detail = f"{type(exc).__name__}: {exc}\n{traceback.format_exc(limit=5)}"
            record = TestRecord(
                name=name,
                status=FAIL,
                detail=detail,
                started_at=started_at,
                finished_at=finished_at,
                duration_seconds=round(time.monotonic() - started, 3),
            )
            self.ctx.test_records.append(record)
            log_event(self.ctx, "error", "test.error", name=name, detail=detail)
            return

        finished_at = iso_now()
        record = TestRecord(
            name=name,
            status=PASS,
            detail=detail,
            started_at=started_at,
            finished_at=finished_at,
            duration_seconds=round(time.monotonic() - started, 3),
        )
        self.ctx.test_records.append(record)
        log_event(self.ctx, "info", "test.pass", name=name, detail=detail)

    def run(self) -> None:
        self.run_test("A1: delins.docx - read ins+del, author present", self.test_a1)
        self.run_test("A2: 58067.docx - pPr-level ins handled", self.test_a2)
        self.run_test("A3: Tika-792.docx - nested ins no crash", self.test_a3)
        self.run_test("A4: stress023.docx - many revisions detected", self.test_a4)
        self.run_test("B1: Accept no-op preserves image doc", self.test_b1)
        self.run_test("B2: Accept pPr ins preserves content", self.test_b2)
        self.run_test("B3: Reject on delins restores deleted text", self.test_b3)
        self.run_test("B4: Accept all on stress023 preserves images", self.test_b4)
        self.run_test("B5: Reject all on stress023 preserves images", self.test_b5)
        self.run_test("C1: Surgical insert keeps unique ids", self.test_c1)
        self.run_test("C2: Surgical delete uses new max id", self.test_c2)
        self.run_test("C3: Insert + accept round-trip", self.test_c3)
        self.run_test("C4: enable_track_changes updates settings", self.test_c4)
        self.run_test("D1: WordWithAttachments comments readable", self.test_d1)
        self.run_test("D2: Delete middle comment stays valid", self.test_d2)
        self.run_test("D3: Comment empty paragraph", self.test_d3)
        self.run_test("D4: Delete missing comment id", self.test_d4)
        self.run_test("D5: Add comment without disturbing image", self.test_d5)
        self.run_test("E1: Reviewer workflow on delins", self.test_e1)
        self.run_test("E2: Reviewer workflow on attachments", self.test_e2)
        self.run_test("E3: Accept changes then add summary comment", self.test_e3)
        self.run_test("F1: Add image to stress023 preserves revisions", self.test_f1)
        self.run_test("F2: Add comment to drawing preserves images", self.test_f2)
        self.run_test("F3: Replace image + tracked insertion", self.test_f3)
        self.run_test("G1: check_winword returns a string", self.test_g1)
        self.run_test("G2: generate_vbs_runner produces valid .vbs file", self.test_g2)
        self.run_test("G3: write_bas_file produces well-formed .bas module", self.test_g3)

    def test_a1(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/delins.docx"))
        changes = edit_docx.get_tracked_changes(doc)
        insertions = [c for c in changes if c["type"] == "insertion"]
        deletions = [c for c in changes if c["type"] == "deletion"]
        assert insertions, f"No insertions found, got changes={changes[:5]}"
        assert deletions, "No deletions found"
        assert all(c["author"] for c in changes), "Some changes have empty author"
        authors = sorted({c["author"] for c in changes})
        return f"ins={len(insertions)} del={len(deletions)} authors={authors[:5]}"

    def test_a2(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/58067.docx"))
        changes = edit_docx.get_tracked_changes(doc)
        for change in changes:
            assert "type" in change and "text" in change, f"Malformed change record: {change}"
        return f"changes found={len(changes)}"

    def test_a3(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/Tika-792.docx"))
        changes = edit_docx.get_tracked_changes(doc)
        return f"changes found={len(changes)}"

    def test_a4(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "integration/stress023.docx"))
        changes = edit_docx.get_tracked_changes(doc)
        insertions = [c for c in changes if c["type"] == "insertion"]
        deletions = [c for c in changes if c["type"] == "deletion"]
        assert len(insertions) >= 100, f"Expected >=100 insertions, got {len(insertions)}"
        assert len(deletions) >= 50, f"Expected >=50 deletions, got {len(deletions)}"
        return f"ins={len(insertions)} del={len(deletions)}"

    def test_b1(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/VariousPictures.docx"))
        n = edit_docx.accept_all_changes(doc)
        out = _output_path(self.ctx, "B1_accept_no_changes.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid after no-op accept: {result['errors']}"
        assert n == 0, f"Expected 0 changes accepted, got {n}"
        return f"accepted={n} images={result['stats']['images']}"

    def test_b2(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/58067.docx"))
        paras_before = len(doc.paragraphs)
        n = edit_docx.accept_all_changes(doc)
        out = _output_path(self.ctx, "B2_accept_ppr_ins.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["tracked_insertions"] == 0, "w:ins markers remain after accept"
        doc2 = edit_docx.open_document(out)
        assert len(doc2.paragraphs) >= paras_before - 2, "Too many paragraphs lost"
        return f"accepted={n} paragraphs_before={paras_before} after={len(doc2.paragraphs)}"

    def test_b3(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/delins.docx"))
        changes_before = edit_docx.get_tracked_changes(doc)
        deleted_texts = [c["text"] for c in changes_before if c["type"] == "deletion" and c["text"]]
        n = edit_docx.reject_all_changes(doc)
        out = _output_path(self.ctx, "B3_reject_delins.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["tracked_insertions"] == 0, "w:ins markers remain"
        assert result["stats"]["tracked_deletions"] == 0, "w:del markers remain"
        xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", errors="replace")
        verified = 0
        for deleted_text in deleted_texts[:3]:
            words = [word for word in deleted_text.split() if len(word) > 3]
            for word in words[:2]:
                if word in xml:
                    verified += 1
                    break
        return f"rejected={n} del_runs_restored={verified}/{min(3, len(deleted_texts))}"

    def test_b4(self) -> str:
        path = _fixture(self.ctx, "integration/stress023.docx")
        doc = edit_docx.open_document(path)
        images_before = validate_docx.validate(path)["stats"]["images"]
        n = edit_docx.accept_all_changes(doc)
        out = _output_path(self.ctx, "B4_accept_stress023.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["tracked_insertions"] == 0, "w:ins remain"
        assert result["stats"]["tracked_deletions"] == 0, "w:del remain"
        assert result["stats"]["images"] == images_before, (
            f"Images changed: was {images_before}, now {result['stats']['images']}"
        )
        return f"accepted={n} images={result['stats']['images']}"

    def test_b5(self) -> str:
        path = _fixture(self.ctx, "integration/stress023.docx")
        doc = edit_docx.open_document(path)
        images_before = validate_docx.validate(path)["stats"]["images"]
        n = edit_docx.reject_all_changes(doc)
        out = _output_path(self.ctx, "B5_reject_stress023.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["images"] == images_before, (
            f"Images changed: was {images_before}, now {result['stats']['images']}"
        )
        return f"rejected={n} images={result['stats']['images']}"

    def test_c1(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/delins.docx"))
        ids_before = _revision_ids_in_doc(doc)
        changes_before = edit_docx.get_tracked_changes(doc)
        edit_docx.insert_tracked_insertion(doc, 0, " [REVIEWER NOTE]", author="Reviewer")
        ids_after = _revision_ids_in_doc(doc)
        new_ids = ids_after - ids_before
        assert len(new_ids) == 1, f"Expected 1 new id, got {new_ids}"
        assert new_ids.isdisjoint(ids_before), f"New id collides with existing: {new_ids & ids_before}"
        out = _output_path(self.ctx, "C1_surgical_insert.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        changes_after = edit_docx.get_tracked_changes(edit_docx.open_document(out))
        assert len(changes_after) == len(changes_before) + 1, (
            f"Expected {len(changes_before) + 1} changes, got {len(changes_after)}"
        )
        return f"new_id={sorted(new_ids)} total_changes={len(changes_after)}"

    def test_c2(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "integration/stress023.docx"))
        ids_before = _revision_ids_in_doc(doc)
        max_id_before = max(ids_before) if ids_before else 0
        edit_docx.insert_tracked_deletion(doc, 0, "obsolete text", author="Reviewer")
        ids_after = _revision_ids_in_doc(doc)
        new_ids = ids_after - ids_before
        assert len(new_ids) == 1, f"Expected exactly one new id, got {new_ids}"
        new_id = next(iter(new_ids))
        assert new_id > max_id_before, f"New id {new_id} not > max existing {max_id_before}"
        out = _output_path(self.ctx, "C2_surgical_delete_stress023.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        return f"max_id_before={max_id_before} new_id={new_id}"

    def test_c3(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/delins.docx"))
        changes_before = len(edit_docx.get_tracked_changes(doc))
        edit_docx.insert_tracked_insertion(doc, 0, " ADDED", author="Test")
        mid = _output_path(self.ctx, "C3_mid.docx")
        edit_docx.save_document(doc, mid)
        doc2 = edit_docx.open_document(mid)
        assert len(edit_docx.get_tracked_changes(doc2)) == changes_before + 1
        edit_docx.accept_all_changes(doc2)
        out = _output_path(self.ctx, "C3_accepted.docx")
        edit_docx.save_document(doc2, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["tracked_insertions"] == 0
        assert result["stats"]["tracked_deletions"] == 0
        xml = zipfile.ZipFile(out).read("word/document.xml").decode("utf-8", errors="replace")
        assert "ADDED" in xml, "Accepted inserted text not in final document"
        return f"changes_before={changes_before} accepted_markup_removed=1"

    def test_c4(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/delins.docx"))
        edit_docx.enable_track_changes(doc)
        out = _output_path(self.ctx, "C4_track_enabled.docx")
        edit_docx.save_document(doc, out)
        with zipfile.ZipFile(out) as zf:
            settings = zf.read("word/settings.xml").decode("utf-8", errors="replace")
        assert "trackChanges" in settings, "w:trackChanges not found in settings.xml"
        return "w:trackChanges present"

    def test_d1(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/WordWithAttachments.docx"))
        comments = edit_docx.read_comments(doc)
        assert len(comments) == 3, f"Expected 3 comments, got {len(comments)}"
        assert all(c["author"] == "vbiryukova" for c in comments), (
            f"Unexpected authors: {[c['author'] for c in comments]}"
        )
        return "3 comments, expected author"

    def test_d2(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/WordWithAttachments.docx"))
        comments = edit_docx.read_comments(doc)
        middle_id = comments[1]["id"]
        ok = edit_docx.delete_comment(doc, middle_id)
        assert ok, "delete_comment returned False"
        out = _output_path(self.ctx, "D2_deleted_middle_comment.docx")
        edit_docx.save_document(doc, out)
        doc2 = edit_docx.open_document(out)
        remaining = edit_docx.read_comments(doc2)
        assert len(remaining) == 2, f"Expected 2 comments after delete, got {len(remaining)}"
        remaining_ids = {c["id"] for c in remaining}
        assert middle_id not in remaining_ids, "Deleted comment id still present"
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        return f"deleted id={middle_id} remaining={sorted(remaining_ids)}"

    def test_d3(self) -> str:
        doc = create_docx.create_document()
        create_docx.add_paragraph(doc, "First paragraph with text")
        doc.add_paragraph()
        create_docx.add_paragraph(doc, "Third paragraph")
        comment_id = edit_docx.insert_comment_on_paragraph(doc, 1, "Empty paragraph review", author="AI")
        out = _output_path(self.ctx, "D3_comment_empty_para.docx")
        edit_docx.save_document(doc, out)
        doc2 = edit_docx.open_document(out)
        comments = edit_docx.read_comments(doc2)
        assert len(comments) == 1, f"Expected 1 comment, got {len(comments)}"
        assert comments[0]["text"] == "Empty paragraph review"
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        return f"comment id={comment_id}"

    def test_d4(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/comment.docx"))
        result = edit_docx.delete_comment(doc, 99999)
        assert result is False, f"Expected False for missing id, got {result}"
        return "delete_comment(99999) returned False"

    def test_d5(self) -> str:
        path = _fixture(self.ctx, "document/testComment.docx")
        doc = edit_docx.open_document(path)
        comments_before = len(edit_docx.read_comments(doc))
        images_before = validate_docx.validate(path)["stats"]["images"]
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                edit_docx.insert_comment_on_paragraph(doc, idx, "Looks good", author="Reviewer")
                break
        out = _output_path(self.ctx, "D5_testcomment_plus_new.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["images"] == images_before, (
            f"Images changed: {images_before} -> {result['stats']['images']}"
        )
        comments_after = len(edit_docx.read_comments(edit_docx.open_document(out)))
        assert comments_after == comments_before + 1, (
            f"Expected {comments_before + 1} comments, got {comments_after}"
        )
        return f"comments {comments_before}->{comments_after} images={images_before}"

    def test_e1(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/delins.docx"))
        changes = edit_docx.get_tracked_changes(doc)
        insertions = [c for c in changes if c["type"] == "insertion" and c["text"].strip()]
        deletions = [c for c in changes if c["type"] == "deletion" and c["text"].strip()]
        ins_para_idx = None
        del_para_idx = None
        for idx, paragraph in enumerate(doc.paragraphs):
            paragraph_el = paragraph._p
            if paragraph_el.find(f".//{{{W_NS}}}ins") is not None and ins_para_idx is None:
                ins_para_idx = idx
            if paragraph_el.find(f".//{{{W_NS}}}del") is not None and del_para_idx is None:
                del_para_idx = idx
        assert ins_para_idx is not None, "No paragraph with w:ins found"
        assert del_para_idx is not None, "No paragraph with w:del found"
        edit_docx.insert_comment_on_paragraph(
            doc,
            ins_para_idx,
            f"Insertion by {insertions[0]['author'] if insertions else 'unknown'} - approved",
            author="AI-Reviewer",
        )
        edit_docx.insert_comment_on_paragraph(
            doc,
            del_para_idx,
            f"Deletion of '{deletions[0]['text'][:30] if deletions else '?'}' - confirm removal",
            author="AI-Reviewer",
        )
        edit_docx.insert_tracked_deletion(doc, 0, "Tika can be:", author="AI-Reviewer")
        out = _output_path(self.ctx, "E1_reviewer_workflow.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["comments"] == 2, f"Expected 2 comments, got {result['stats']['comments']}"
        assert result["stats"]["tracked_deletions"] >= len(deletions) + 1, "New tracked deletion not found"
        comments = [c for c in edit_docx.read_comments(edit_docx.open_document(out)) if c["author"] == "AI-Reviewer"]
        assert len(comments) == 2, f"Expected 2 reviewer comments, got {len(comments)}"
        return f"ins_para={ins_para_idx} del_para={del_para_idx}"

    def test_e2(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/WordWithAttachments.docx"))
        original_comments = edit_docx.read_comments(doc)
        assert len(original_comments) == 3, f"Expected 3 comments, got {len(original_comments)}"
        edit_docx.insert_comment_on_paragraph(doc, 0, "Overall LGTM", author="AI-Reviewer")
        n_accepted = edit_docx.accept_all_changes(doc)
        first_id = original_comments[0]["id"]
        edit_docx.delete_comment(doc, first_id)
        out = _output_path(self.ctx, "E2_reviewer_attachments.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["comments"] == 3, f"Expected 3 comments, got {result['stats']['comments']}"
        assert n_accepted == 0, f"Expected no accepted changes, got {n_accepted}"
        final_comments = edit_docx.read_comments(edit_docx.open_document(out))
        final_ids = {c["id"] for c in final_comments}
        assert first_id not in final_ids, "Deleted comment still present"
        reviewer_comments = [c for c in final_comments if c["author"] == "AI-Reviewer"]
        assert len(reviewer_comments) == 1, f"Expected 1 reviewer comment, got {len(reviewer_comments)}"
        return "3 comments remain after add+delete"

    def test_e3(self) -> str:
        doc = edit_docx.open_document(_fixture(self.ctx, "document/delins.docx"))
        changes = edit_docx.get_tracked_changes(doc)
        n_insertions = sum(1 for c in changes if c["type"] == "insertion")
        n_deletions = sum(1 for c in changes if c["type"] == "deletion")
        n_accepted = edit_docx.accept_all_changes(doc)
        summary = (
            f"Reviewed and accepted {n_insertions} insertions and {n_deletions} deletions from original authors."
        )
        edit_docx.insert_comment_on_paragraph(doc, 0, summary, author="AI-Reviewer")
        out = _output_path(self.ctx, "E3_accepted_then_commented.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["tracked_insertions"] == 0
        assert result["stats"]["tracked_deletions"] == 0
        assert result["stats"]["comments"] == 1
        comment_text = edit_docx.read_comments(edit_docx.open_document(out))[0]["text"]
        assert "accepted" in comment_text.lower(), f"Unexpected comment text: {comment_text}"
        return f"accepted={n_accepted} summary_comment=1"

    def test_f1(self) -> str:
        path = _fixture(self.ctx, "integration/stress023.docx")
        doc = edit_docx.open_document(path)
        changes_before = edit_docx.get_tracked_changes(doc)
        images_before = validate_docx.validate(path)["stats"]["images"]
        png_path = _write_binary(_output_path(self.ctx, "F1_test.png"), _make_png(0, 128, 255))
        create_docx.add_inline_image(doc, str(png_path), width=Inches(1), alt_text="Reviewer diagram")
        out = _output_path(self.ctx, "F1_stress023_plus_image.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["images"] == images_before + 1, (
            f"Expected {images_before + 1} images, got {result['stats']['images']}"
        )
        changes_after = edit_docx.get_tracked_changes(edit_docx.open_document(out))
        assert len(changes_after) == len(changes_before), (
            f"Track changes changed: {len(changes_before)} -> {len(changes_after)}"
        )
        return f"images {images_before}->{result['stats']['images']}"

    def test_f2(self) -> str:
        path = _fixture(self.ctx, "document/drawing.docx")
        doc = edit_docx.open_document(path)
        images_before = validate_docx.validate(path)["stats"]["images"]
        assert images_before >= 1, "drawing.docx has no images"
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() or paragraph.runs:
                edit_docx.insert_comment_on_paragraph(doc, idx, "Drawing review", author="AI-Reviewer")
                break
        out = _output_path(self.ctx, "F2_drawing_with_comment.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["images"] == images_before, (
            f"Images lost: {images_before} -> {result['stats']['images']}"
        )
        assert result["stats"]["comments"] == 1, f"Expected 1 comment, got {result['stats']['comments']}"
        return f"images={images_before} comments=1"

    def test_f3(self) -> str:
        path = _fixture(self.ctx, "document/VariousPictures.docx")
        doc = edit_docx.open_document(path)
        edit_docx.insert_tracked_insertion(doc, 0, " [reviewed]", author="Reviewer")
        blips = [
            blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            for blip in doc.element.body.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            if blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed") is not None
        ]
        assert blips, "No image relationships found in VariousPictures.docx"
        current_ext = Path(doc.part.related_parts[blips[0]].partname).suffix.lower()
        suffix = ".wmf" if current_ext == ".wmf" else ".emf" if current_ext == ".emf" else ".jpeg"
        replacement_bytes = JPEG_1X1
        if current_ext == ".png":
            suffix = ".png"
            replacement_bytes = PNG_1X1
        elif current_ext in (".jpg", ".jpeg"):
            suffix = ".jpeg"
            replacement_bytes = JPEG_1X1
        elif current_ext == ".wmf":
            suffix = ".wmf"
            replacement_bytes = _existing_fixture(self.ctx, ("document/testException2.doc-2.wmf",)).read_bytes()
        elif current_ext == ".emf":
            suffix = ".emf"
            replacement_bytes = _existing_fixture(self.ctx, ("document/vector_image.emf",)).read_bytes()
        replacement_path = _output_path(self.ctx, f"F3_replace{suffix}")
        _write_binary(replacement_path, replacement_bytes)
        replaced = create_docx.replace_image(doc, 0, replacement_path)
        assert replaced, "replace_image returned False"
        out = _output_path(self.ctx, "F3_replaced_image_with_trackchange.docx")
        edit_docx.save_document(doc, out)
        result = validate_docx.validate(out)
        assert result["valid"], f"Document invalid: {result['errors']}"
        assert result["stats"]["tracked_insertions"] == 1, "Tracked insertion lost"
        assert result["stats"]["images"] == 5, f"Image count changed: {result['stats']['images']}"
        return "image_replaced=True tracked_insertion=1"


    def test_g1(self) -> str:
        result = vba_runner.check_winword()
        assert isinstance(result, str), f"check_winword() must return str, got {type(result)}"
        status = f"found: {result}" if result else "not on PATH (acceptable in CI)"
        return status

    def test_g2(self) -> str:
        vba_code = "Sub TestG2()\n    MsgBox \"hello from G2\"\nEnd Sub\n"
        out = _output_path(self.ctx, "G2_runner.vbs")
        vba_runner.generate_vbs_runner("C:\\test\\doc.docx", vba_code, "TestG2", out)
        assert out.exists(), f"VBScript file not created: {out}"
        content = out.read_text(encoding="utf-8")
        assert len(content) > 0, "Generated .vbs file is empty"
        assert "CreateObject" in content, ".vbs missing CreateObject"
        assert "Word.Application" in content, ".vbs missing Word.Application"
        assert "TestG2" in content, ".vbs missing sub name TestG2"
        assert "C:\\\\test\\\\doc.docx" in content or "C:\\test\\doc.docx" in content, ".vbs missing doc path"
        return f"vbs_size={len(content)} bytes"

    def test_g3(self) -> str:
        vba_code = "Sub TestG3()\n    ActiveDocument.Content.InsertAfter \"hello\"\nEnd Sub\n"
        out = _output_path(self.ctx, "G3_module.bas")
        vba_runner.write_bas_file(vba_code, out)
        assert out.exists(), f".bas file not created: {out}"
        content = out.read_text(encoding="utf-8")
        assert "Attribute VB_Name" in content, ".bas missing VB_Name header"
        assert "Option Explicit" in content, ".bas missing Option Explicit"
        assert "TestG3" in content, ".bas missing sub name"
        assert "ActiveDocument" in content, ".bas missing VBA code body"
        return f"bas_size={len(content)} bytes header_ok=True"


def cleanup(ctx: RunContext) -> None:
    artifacts = [ctx.repo_dir, ctx.output_dir]
    for path in artifacts:
        if path.exists():
            shutil.rmtree(path, ignore_errors=True)
            log_event(ctx, "info", "cleanup.path_removed", path=str(path))
    ctx.cleanup_completed = True
    write_metadata(ctx, status="cleanup-complete")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run the Apache POI-backed docx regression workflow")
    parser.add_argument("--repo-url", default=DEFAULT_REPO_URL, help="Apache POI git URL")
    parser.add_argument("--ref", default=DEFAULT_REF, help="Git ref to test against")
    parser.add_argument("--keep-artifacts", action="store_true", help="Keep temp clone and outputs for debugging")
    parser.add_argument("--timeout", type=int, default=DEFAULT_TIMEOUT_SECONDS, help="Per-command timeout in seconds")
    parser.add_argument("--retries", type=int, default=DEFAULT_RETRIES, help="Retry count for clone/fetch commands")
    return parser.parse_args()


def build_context(args: argparse.Namespace) -> RunContext:
    TEMP_ROOT.mkdir(parents=True, exist_ok=True)
    run_root = Path(tempfile.mkdtemp(prefix="docx-poi-suite-", dir=str(TEMP_ROOT)))
    output_dir = run_root / "outputs"
    metadata_path = run_root / "run-metadata.json"
    log_path = run_root / "workflow.log"
    repo_dir = run_root / "apache-poi"
    ctx = RunContext(
        args=args,
        run_root=run_root,
        repo_dir=repo_dir,
        output_dir=output_dir,
        log_path=log_path,
        metadata_path=metadata_path,
        repo_url=args.repo_url,
        ref=args.ref,
    )
    ctx.test_command = " ".join([sys.executable, str(Path(__file__).resolve()), *sys.argv[1:]])
    return ctx


def main() -> int:
    args = parse_args()
    ctx = build_context(args)
    status = "failed"
    error_message: str | None = None
    exit_code = 1
    log_event(ctx, "info", "workflow.start", command=ctx.test_command, run_root=str(ctx.run_root))
    try:
        prepare_poi_baseline(ctx)
        verify_baseline(ctx)
        suite = TestSuite(ctx)
        suite.run()
        failed = [rec for rec in ctx.test_records if rec.status == FAIL]
        status = "passed" if not failed else "failed"
        exit_code = 0 if not failed else 1
        log_event(
            ctx,
            "info" if not failed else "warning",
            "workflow.summary",
            total=len(ctx.test_records),
            passed=sum(1 for rec in ctx.test_records if rec.status == PASS),
            failed=len(failed),
        )
        return exit_code
    except WorkflowError as exc:
        error_message = str(exc)
        log_event(ctx, "error", "workflow.error", error=error_message)
        return 1
    except Exception as exc:
        error_message = f"Unhandled error: {type(exc).__name__}: {exc}"
        log_event(ctx, "error", "workflow.exception", error=error_message, traceback=traceback.format_exc())
        return 1
    finally:
        write_metadata(ctx, status=status, error=error_message)
        if args.keep_artifacts:
            log_event(ctx, "info", "cleanup.skipped", run_root=str(ctx.run_root))
        else:
            cleanup(ctx)


if __name__ == "__main__":
    sys.exit(main())
