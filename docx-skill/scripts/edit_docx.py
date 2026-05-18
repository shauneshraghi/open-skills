"""
edit_docx.py — Edit existing .docx files: comments and track changes.

Design decisions sourced from:
- python-docx 1.x API: Document.add_comment(), Document.comments (native since 1.0)
- ECMA-376 §17.13.4: comment XML structure (commentRangeStart/End, commentReference)
- ECMA-376 §17.13.5: tracked change elements (w:ins, w:del, w:delText)
- python-docx oxml/comments.py: CT_Comments.add_comment() shows the id-generation algorithm
- python-docx oxml/text/run.py: insert_comment_range_start_above() / *_end_and_reference_below()

NOTE: Comment deletion and all track-change operations require direct lxml
manipulation — python-docx 1.2 provides no API for these.
"""

from __future__ import annotations

import datetime
from pathlib import Path
from typing import Union

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

# ── Namespace URIs ────────────────────────────────────────────────────────────
_W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML = "http://www.w3.org/XML/1998/namespace"


# ── Open / Save ───────────────────────────────────────────────────────────────


def open_document(path: Union[str, Path]) -> Document:
    return Document(str(path))


def save_document(doc: Document, path: Union[str, Path]) -> None:
    doc.save(str(path))


# ── Comments ──────────────────────────────────────────────────────────────────


def insert_comment(
    doc: Document,
    run: Run,
    text: str,
    author: str = "",
    initials: str = "",
) -> int:
    """Attach a comment to `run` and return the new comment id.

    Uses the python-docx 1.x native API (Document.add_comment).
    Internally python-docx inserts:
      - w:commentRangeStart before the run
      - w:commentRangeEnd + w:commentReference run after the run
      - w:comment element in comments.xml

    Source: python-docx docs — Document.add_comment(runs, text, author, initials)
    """
    comment = doc.add_comment(run, text=text, author=author, initials=initials)
    return comment.comment_id


def insert_comment_on_paragraph(
    doc: Document,
    paragraph_index: int,
    text: str,
    author: str = "",
    initials: str = "",
) -> int:
    """Convenience: comment on all runs in paragraph at `paragraph_index`.

    If the paragraph has no runs, adds a placeholder run first.
    Returns the new comment id.
    """
    para = doc.paragraphs[paragraph_index]
    runs = list(para.runs)
    if not runs:
        run = para.add_run(" ")
        runs = [run]
    comment = doc.add_comment(runs, text=text, author=author, initials=initials)
    return comment.comment_id


def read_comments(doc: Document) -> list[dict]:
    """Return all comments as a list of dicts with keys:
    id, author, initials, timestamp (ISO string or None), text.

    Source: python-docx Comments.__iter__() yields Comment objects with these properties.
    """
    results = []
    for c in doc.comments:
        results.append({
            "id":        c.comment_id,
            "author":    c.author,
            "initials":  c.initials,
            "timestamp": c.timestamp.isoformat() if c.timestamp else None,
            "text":      c.text,
        })
    return results


def delete_comment(doc: Document, comment_id: int) -> bool:
    """Delete the comment with `comment_id` and all its document markers.

    python-docx has no native deletion API, so we manipulate lxml directly.

    Per ECMA-376 §17.13.4, a comment is anchored by four XML elements:
      1. w:commentRangeStart w:id="N"  — before the commented content
      2. w:commentRangeEnd   w:id="N"  — after the commented content
      3. w:commentReference  w:id="N"  — inside a w:r after the range end
      4. w:comment           w:id="N"  — in word/comments.xml

    All four must be removed to leave a well-formed document.
    Returns True if the comment was found and removed.
    """
    cid = str(comment_id)
    body = doc.element.body
    found = False

    # Remove w:commentRangeStart and w:commentRangeEnd from document body
    for tag in ("w:commentRangeStart", "w:commentRangeEnd"):
        for el in list(body.iter(qn(tag))):
            if el.get(qn("w:id")) == cid:
                el.getparent().remove(el)
                found = True

    # Remove the w:r that contains the w:commentReference
    for ref_el in list(body.iter(qn("w:commentReference"))):
        if ref_el.get(qn("w:id")) == cid:
            run_el = ref_el.getparent()
            if run_el is not None:
                run_parent = run_el.getparent()
                if run_parent is not None:
                    run_parent.remove(run_el)
            found = True

    # Remove the w:comment element from comments.xml.
    # python-docx 1.x exposes the comments part as doc.part._comments_part (private).
    # The public doc.comments._comments_elm is an equivalent route that avoids the
    # private attribute. We try both so the code works across minor version changes.
    comments_elm = None
    try:
        comments_elm = doc.part._comments_part.element
    except AttributeError:
        pass
    if comments_elm is None:
        try:
            comments_elm = doc.comments._comments_elm
        except AttributeError:
            pass
    if comments_elm is not None:
        for c_el in list(comments_elm.findall(qn("w:comment"))):
            if c_el.get(qn("w:id")) == cid:
                comments_elm.remove(c_el)
                found = True

    return found


def delete_all_comments(doc: Document) -> int:
    """Delete every comment in the document. Returns the count removed."""
    ids = [c.comment_id for c in doc.comments]
    for cid in ids:
        delete_comment(doc, cid)
    return len(ids)


# ── Track Changes ─────────────────────────────────────────────────────────────
# All track-change operations require direct lxml manipulation.
# Source: ECMA-376 §17.13.5 (revision markup), §17.13.5.16 (w:ins), §17.13.5.14 (w:del)


def _next_revision_id(doc: Document) -> int:
    """Return the next unique revision id for w:ins / w:del elements.

    Scans all w:ins, w:del, and w:commentRangeStart elements for their w:id
    values and returns max + 1.  The id space is shared across all revision
    markup per ECMA-376 §17.13.5.
    """
    body = doc.element.body
    used = set()
    for tag in ("w:ins", "w:del", "w:commentRangeStart"):
        for el in body.iter(qn(tag)):
            try:
                used.add(int(el.get(qn("w:id"), "0")))
            except ValueError:
                pass
    return (max(used) + 1) if used else 1


def _utc_now_str() -> str:
    return datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")


def insert_tracked_insertion(
    doc: Document,
    paragraph_index: int,
    text: str,
    author: str = "",
    position: str = "end",
) -> None:
    """Insert text as a tracked insertion (w:ins) in the paragraph.

    The w:ins element wraps a w:r containing a w:t.  When changes are accepted,
    the run content is kept; when rejected, it is discarded.

    Args:
        position: 'end' to append to the paragraph, 'start' to prepend.

    Source: ECMA-376 §17.13.5.16 — w:ins is a run-level revision element.
    Direct lxml is required; python-docx has no track-change API.
    """
    para = doc.paragraphs[paragraph_index]
    p_el = para._p
    rev_id = _next_revision_id(doc)

    ins_el = etree.SubElement(p_el, f"{{{_W}}}ins") if position == "end" else None
    if ins_el is None:
        # prepend: insert before the first child
        ins_el = etree.Element(f"{{{_W}}}ins")
        p_el.insert(0, ins_el)

    ins_el.set(f"{{{_W}}}id",     str(rev_id))
    ins_el.set(f"{{{_W}}}author", author)
    ins_el.set(f"{{{_W}}}date",   _utc_now_str())

    r_el = etree.SubElement(ins_el, f"{{{_W}}}r")
    t_el = etree.SubElement(r_el,   f"{{{_W}}}t")
    t_el.text = text
    # xml:space="preserve" prevents stripping of leading/trailing spaces
    # Source: ECMA-376 §22.9.2.19 — XML space preservation
    t_el.set(f"{{{_XML}}}space", "preserve")


def insert_tracked_deletion(
    doc: Document,
    paragraph_index: int,
    text: str,
    author: str = "",
) -> None:
    """Insert text as a tracked deletion (w:del) in the paragraph.

    The w:del wraps a w:r that uses w:delText (not w:t).  When accepted, the
    content is discarded; when rejected, it is restored.

    Source: ECMA-376 §17.13.5.14 — w:del uses w:delText descendants.
    Direct lxml required; python-docx has no track-change API.
    """
    para = doc.paragraphs[paragraph_index]
    p_el = para._p
    rev_id = _next_revision_id(doc)

    del_el = etree.SubElement(p_el, f"{{{_W}}}del")
    del_el.set(f"{{{_W}}}id",     str(rev_id))
    del_el.set(f"{{{_W}}}author", author)
    del_el.set(f"{{{_W}}}date",   _utc_now_str())

    r_el  = etree.SubElement(del_el, f"{{{_W}}}r")
    dt_el = etree.SubElement(r_el,   f"{{{_W}}}delText")
    dt_el.text = text
    dt_el.set(f"{{{_XML}}}space", "preserve")


def accept_all_changes(doc: Document) -> int:
    """Accept all tracked insertions and deletions. Returns count of changes processed.

    Accept algorithm (ECMA-376 §17.13.5 semantics):
    - w:ins: replace the element with its w:r children (keep inserted content)
    - w:del: remove the element and all children (discard deleted content)

    We must collect elements before iterating because we modify the tree.
    """
    body = doc.element.body
    count = 0

    # Accept insertions: lift runs out of w:ins, remove the w:ins wrapper
    for ins_el in list(body.iter(f"{{{_W}}}ins")):
        parent = ins_el.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins_el)
        for child in list(ins_el):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins_el)
        count += 1

    # Accept deletions: remove the w:del and all its content
    for del_el in list(body.iter(f"{{{_W}}}del")):
        parent = del_el.getparent()
        if parent is not None:
            parent.remove(del_el)
            count += 1

    return count


def reject_all_changes(doc: Document) -> int:
    """Reject all tracked insertions and deletions. Returns count of changes processed.

    Reject algorithm (ECMA-376 §17.13.5 semantics):
    - w:ins: remove the element and all children (discard inserted content)
    - w:del: restore runs by converting w:delText → w:t and lifting them out

    We must convert w:delText to w:t because the w:t element name is what
    python-docx and Word use for run text content in non-deleted context.
    Source: ECMA-376 §17.3.3.8 (w:delText) vs §17.3.3.32 (w:t).
    """
    body = doc.element.body
    count = 0

    # Reject deletions: restore the original text by converting delText→t
    # Also convert delInstrText→instrText to restore deleted hyperlink/field instructions.
    # Source: ECMA-376 §17.3.3.10 (w:delInstrText) vs §17.3.3.23 (w:instrText)
    for del_el in list(body.iter(f"{{{_W}}}del")):
        parent = del_el.getparent()
        if parent is None:
            continue
        idx = list(parent).index(del_el)
        for r_el in list(del_el):
            for del_text in r_el.findall(f"{{{_W}}}delText"):
                del_text.tag = f"{{{_W}}}t"
            for del_instr in r_el.findall(f"{{{_W}}}delInstrText"):
                del_instr.tag = f"{{{_W}}}instrText"
            parent.insert(idx, r_el)
            idx += 1
        parent.remove(del_el)
        count += 1

    # Reject insertions: discard the inserted content
    for ins_el in list(body.iter(f"{{{_W}}}ins")):
        parent = ins_el.getparent()
        if parent is not None:
            parent.remove(ins_el)
            count += 1

    return count


def enable_track_changes(doc: Document) -> None:
    """Set w:trackChanges in settings.xml so Word continues tracking on open.

    Source: ECMA-376 §17.15.1.89 — w:trackChanges element in w:settings.
    python-docx exposes settings via doc.settings.element.
    """
    settings_el = doc.settings.element
    tag = f"{{{_W}}}trackChanges"
    if settings_el.find(tag) is None:
        tc_el = etree.SubElement(settings_el, tag)
        # Insert at position 0 (Word convention: trackChanges near the top)
        settings_el.remove(tc_el)
        settings_el.insert(0, tc_el)


def get_tracked_changes(doc: Document) -> list[dict]:
    """Return a summary of all tracked insertions and deletions.

    Returns a list of dicts: {type: 'insertion'|'deletion', author, date, text}
    """
    body = doc.element.body
    results = []

    for ins_el in body.iter(f"{{{_W}}}ins"):
        text_parts = [
            t.text or ""
            for t in ins_el.iter(f"{{{_W}}}t")
        ]
        results.append({
            "type":   "insertion",
            "author": ins_el.get(f"{{{_W}}}author", ""),
            "date":   ins_el.get(f"{{{_W}}}date", ""),
            "text":   "".join(text_parts),
        })

    for del_el in body.iter(f"{{{_W}}}del"):
        text_parts = [
            t.text or ""
            for t in del_el.iter(f"{{{_W}}}delText")
        ]
        results.append({
            "type":   "deletion",
            "author": del_el.get(f"{{{_W}}}author", ""),
            "date":   del_el.get(f"{{{_W}}}date", ""),
            "text":   "".join(text_parts),
        })

    return results


# ── CLI ───────────────────────────────────────────────────────────────────────


def _cli() -> None:
    import argparse
    import json
    import sys

    parser = argparse.ArgumentParser(description="Edit a .docx file")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_comment = sub.add_parser("add-comment", help="Add a comment to a paragraph")
    p_comment.add_argument("input")
    p_comment.add_argument("output")
    p_comment.add_argument("--para", type=int, default=0)
    p_comment.add_argument("--text", required=True)
    p_comment.add_argument("--author", default="")

    p_read = sub.add_parser("read-comments", help="Read all comments")
    p_read.add_argument("input")

    p_del_comment = sub.add_parser("delete-comment", help="Delete a comment by id")
    p_del_comment.add_argument("input")
    p_del_comment.add_argument("output")
    p_del_comment.add_argument("--id", type=int, required=True)

    p_ins = sub.add_parser("track-insert", help="Insert tracked insertion")
    p_ins.add_argument("input")
    p_ins.add_argument("output")
    p_ins.add_argument("--para", type=int, default=0)
    p_ins.add_argument("--text", required=True)
    p_ins.add_argument("--author", default="")

    p_del = sub.add_parser("track-delete", help="Insert tracked deletion")
    p_del.add_argument("input")
    p_del.add_argument("output")
    p_del.add_argument("--para", type=int, default=0)
    p_del.add_argument("--text", required=True)
    p_del.add_argument("--author", default="")

    p_accept = sub.add_parser("accept", help="Accept all tracked changes")
    p_accept.add_argument("input")
    p_accept.add_argument("output")

    p_reject = sub.add_parser("reject", help="Reject all tracked changes")
    p_reject.add_argument("input")
    p_reject.add_argument("output")

    args = parser.parse_args()

    if args.cmd == "read-comments":
        doc = open_document(args.input)
        comments = read_comments(doc)
        print(json.dumps(comments, indent=2))
        return

    doc = open_document(args.input)

    if args.cmd == "add-comment":
        cid = insert_comment_on_paragraph(doc, args.para, args.text, args.author)
        print(f"Added comment id={cid}", file=sys.stderr)

    elif args.cmd == "delete-comment":
        ok = delete_comment(doc, args.id)
        print(f"Deleted: {ok}", file=sys.stderr)

    elif args.cmd == "track-insert":
        insert_tracked_insertion(doc, args.para, args.text, args.author)
        print("Inserted tracked insertion", file=sys.stderr)

    elif args.cmd == "track-delete":
        insert_tracked_deletion(doc, args.para, args.text, args.author)
        print("Inserted tracked deletion", file=sys.stderr)

    elif args.cmd == "accept":
        n = accept_all_changes(doc)
        print(f"Accepted {n} changes", file=sys.stderr)

    elif args.cmd == "reject":
        n = reject_all_changes(doc)
        print(f"Rejected {n} changes", file=sys.stderr)

    save_document(doc, args.output)
    print(f"Saved: {args.output}", file=sys.stderr)


if __name__ == "__main__":
    _cli()
