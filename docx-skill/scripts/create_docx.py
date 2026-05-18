"""
create_docx.py — Create .docx files with text, tables, inline images, and anchored images.

Design decisions sourced from:
- python-docx docs: add_picture() wraps CT_Inline.new_pic_inline() which builds wp:inline XML
- ECMA-376 §20.4: wp:anchor structure for floating images; EMU units (1 in = 914400)
- ECMA-376 §20.4.2.5: wp:docPr/@descr carries the accessibility alt-text attribute
- python-docx source (oxml/shape.py): CT_Inline._inline_xml() template, CT_Picture.new()
"""

from __future__ import annotations

import copy
import io
from pathlib import Path
from typing import IO, Union

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.shared import Emu, Inches, Pt
from lxml import etree

# ── Namespace URIs ────────────────────────────────────────────────────────────
# Source: python-docx oxml/ns.py nsmap dict; ECMA-376 Annex A namespace table
_W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_XML = "http://www.w3.org/XML/1998/namespace"


def create_document(
    title: str = "",
    author: str = "",
    subject: str = "",
) -> Document:
    """Return a new blank Document with optional core properties set."""
    doc = Document()
    props = doc.core_properties
    if title:
        props.title = title
    if author:
        props.author = author
    if subject:
        props.subject = subject
    return doc


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading paragraph at the given outline level (1–9)."""
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    """Add a body paragraph with optional named style."""
    doc.add_paragraph(text, style=style)


def add_table(
    doc: Document,
    data: list[list[str]],
    style: str = "Table Grid",
) -> None:
    """Add a table populated from a list-of-rows list-of-cells."""
    if not data:
        return
    rows, cols = len(data), max(len(row) for row in data)
    table = doc.add_table(rows=rows, cols=cols, style=style)
    for r, row in enumerate(data):
        for c, text in enumerate(row):
            table.cell(r, c).text = text


# ── Inline images ─────────────────────────────────────────────────────────────


def add_inline_image(
    doc: Document,
    image: Union[str, Path, IO[bytes]],
    width: Emu | None = None,
    height: Emu | None = None,
    alt_text: str = "",
    paragraph_index: int | None = None,
) -> None:
    """Add an inline image to the document.

    The image flows with text (like a large character).  Width/height may be
    omitted to use the image's native DPI dimensions.

    Alt text is written to wp:docPr/@descr per ECMA-376 §20.4.2.5.

    Design: python-docx add_picture() delegates to BlockItemContainer.add_picture()
    which calls CT_Inline.new_pic_inline() to build the wp:inline XML.  We
    post-process the resulting element to inject the descr attribute.
    """
    if paragraph_index is not None and paragraph_index < len(doc.paragraphs):
        para = doc.paragraphs[paragraph_index]
        run = para.add_run()
    else:
        # add_picture appends to the last paragraph (or a new one)
        para = doc.add_paragraph()
        run = para.add_run()

    run.add_picture(image, width=width, height=height)

    if alt_text:
        _set_inline_alt_text(run, alt_text)


def _set_inline_alt_text(run, alt_text: str) -> None:
    """Set descr attribute on wp:docPr inside the run's inline drawing.

    Source: ECMA-376 §20.4.2.5 — CT_NonVisualDrawingProps/@descr carries the
    alternative text string displayed by accessibility tools.
    """
    drawing = run._r.find(qn("w:drawing"))
    if drawing is None:
        return
    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        return
    docPr = inline.find(qn("wp:docPr"))
    if docPr is not None:
        docPr.set("descr", alt_text)


# ── Anchored (floating) images ────────────────────────────────────────────────


def add_anchored_image(
    doc: Document,
    image: Union[str, Path, IO[bytes]],
    width: Emu = Inches(2),
    height: Emu | None = None,
    alt_text: str = "",
    h_align: str = "left",
    v_align: str = "top",
    wrap: str = "square",
    behind_doc: bool = False,
) -> None:
    """Add a floating (anchored) image to the document.

    python-docx has no native API for anchored images — this is documented in
    the python-docx issue tracker as out of scope for the high-level API.
    We build the wp:anchor XML manually per ECMA-376 §20.4.

    Args:
        h_align: Horizontal alignment. One of 'left', 'center', 'right'.
        v_align: Vertical alignment. One of 'top', 'center', 'bottom'.
        wrap:    Text wrap mode. One of 'square', 'tight', 'through', 'none'.
        behind_doc: If True, the image renders behind document text.
    """
    # Add image as a relationship and get back its rId + dimensions
    if isinstance(image, (str, Path)):
        image_path = str(image)
        with open(image_path, "rb") as f:
            image_bytes = f.read()
        image_stream = io.BytesIO(image_bytes)
    else:
        image_bytes = image.read()
        image_stream = io.BytesIO(image_bytes)

    # Add a throwaway inline image first so python-docx registers the image
    # part and gives us the rId.  We will replace the wp:inline with wp:anchor.
    para = doc.add_paragraph()
    run = para.add_run()
    image_stream.seek(0)
    run.add_picture(image_stream, width=width, height=height)

    drawing_el = run._r.find(qn("w:drawing"))
    inline_el  = drawing_el.find(qn("wp:inline"))
    extent_el  = inline_el.find(qn("wp:extent"))
    cx = int(extent_el.get("cx"))
    cy = int(extent_el.get("cy"))
    docPr_el = inline_el.find(qn("wp:docPr"))
    shape_id = int(docPr_el.get("id"))
    shape_name = docPr_el.get("name", f"Picture {shape_id}")

    # Extract the a:graphic subtree to reuse in the anchor
    graphic_el = copy.deepcopy(inline_el.find(qn("a:graphic")))

    # Build the wp:anchor element
    anchor_el = _build_anchor(
        cx=cx, cy=cy,
        shape_id=shape_id, shape_name=shape_name,
        alt_text=alt_text,
        h_align=h_align, v_align=v_align,
        wrap=wrap, behind_doc=behind_doc,
        graphic_el=graphic_el,
    )

    # Swap inline → anchor in the w:drawing element
    drawing_el.remove(inline_el)
    drawing_el.append(anchor_el)


def _build_anchor(
    cx: int, cy: int,
    shape_id: int, shape_name: str,
    alt_text: str,
    h_align: str, v_align: str,
    wrap: str, behind_doc: bool,
    graphic_el,
) -> etree._Element:
    """Build a wp:anchor element per ECMA-376 §20.4.2.3.

    relativeHeight=251658240 is the default z-order used by Word (level 1 above text).
    distL/distR=114300 EMU ≈ 0.125 inch, the default image margin.
    """
    behind = "1" if behind_doc else "0"
    wrap_tag = {
        "square":  "wp:wrapSquare",
        "tight":   "wp:wrapTight",
        "through": "wp:wrapThrough",
        "none":    "wp:wrapNone",
    }.get(wrap, "wp:wrapSquare")

    ns = f'xmlns:wp="{_WP}" xmlns:a="{_A}" xmlns:pic="{_PIC}" xmlns:r="{_R}"'
    wrap_xml = (
        f'<{wrap_tag} wrapText="bothSides" xmlns:wp="{_WP}"/>'
        if wrap != "none"
        else f'<{wrap_tag} xmlns:wp="{_WP}"/>'
    )

    anchor_xml = (
        f'<wp:anchor {ns} '
        f'distT="0" distB="0" distL="114300" distR="114300" '
        f'simplePos="0" relativeHeight="251658240" behindDoc="{behind}" '
        f'locked="0" layoutInCell="1" allowOverlap="1">'
        f'  <wp:simplePos x="0" y="0"/>'
        f'  <wp:positionH relativeFrom="column">'
        f'    <wp:align>{h_align}</wp:align>'
        f'  </wp:positionH>'
        f'  <wp:positionV relativeFrom="paragraph">'
        f'    <wp:align>{v_align}</wp:align>'
        f'  </wp:positionV>'
        f'  <wp:extent cx="{cx}" cy="{cy}"/>'
        f'  <wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'  {wrap_xml}'
        f'  <wp:docPr id="{shape_id}" name="{shape_name}" descr="{alt_text}"/>'
        f'  <wp:cNvGraphicFramePr/>'
        f'</wp:anchor>'
    )
    anchor = parse_xml(anchor_xml)
    anchor.append(graphic_el)
    return anchor


# ── Image replacement ─────────────────────────────────────────────────────────


def replace_image(
    doc: Document,
    image_index: int,
    new_image: Union[str, Path, IO[bytes]],
) -> bool:
    """Replace the image at position `image_index` (0-based) in the document.

    Finds the nth a:blip element and swaps the image bytes in the referenced part.
    The image dimensions in the XML are NOT updated — the new image is scaled to
    the existing frame size.

    Returns True if the replacement was made, False if the index was out of range.

    Design: Relationship-based image parts (ECMA-376 §15.3) — the rId on a:blip
    points to a part in word/media/.  python-docx exposes the part via
    doc.part.related_parts[rId] and its blob via ._blob.
    """
    if isinstance(new_image, (str, Path)):
        with open(str(new_image), "rb") as f:
            new_bytes = f.read()
    else:
        new_bytes = new_image.read()

    blips = doc.element.body.findall(
        f".//{{{_A}}}blip"
    )
    if image_index >= len(blips):
        return False

    blip = blips[image_index]
    rid = blip.get(f"{{{_R}}}embed")
    if rid is None:
        return False

    image_part = doc.part.related_parts[rid]
    image_part._blob = new_bytes
    return True


# ── Convenience: save ─────────────────────────────────────────────────────────


def save(doc: Document, path: Union[str, Path]) -> None:
    """Save the document to disk."""
    doc.save(str(path))


# ── CLI ───────────────────────────────────────────────────────────────────────


def _cli() -> None:
    import argparse
    import sys

    parser = argparse.ArgumentParser(description="Create a sample .docx file")
    parser.add_argument("output", help="Output .docx path")
    parser.add_argument("--title", default="Sample Document")
    parser.add_argument("--author", default="")
    parser.add_argument("--inline-image", metavar="PATH", help="Embed inline image")
    parser.add_argument("--anchored-image", metavar="PATH", help="Embed anchored image")
    parser.add_argument("--alt-text", default="", help="Alt text for image")
    args = parser.parse_args()

    doc = create_document(title=args.title, author=args.author)
    add_heading(doc, args.title, level=1)
    add_paragraph(doc, "This document was created by the docx-creation-editing skill.")
    add_table(doc, [["Name", "Value"], ["Feature", "Images + Comments + Track Changes"]])

    if args.inline_image:
        add_paragraph(doc, "Inline image below:")
        add_inline_image(doc, args.inline_image, width=Inches(3), alt_text=args.alt_text)

    if args.anchored_image:
        add_paragraph(doc, "Anchored image:")
        add_anchored_image(doc, args.anchored_image, width=Inches(2), alt_text=args.alt_text)

    save(doc, args.output)
    print(f"Saved: {args.output}", file=sys.stderr)


if __name__ == "__main__":
    _cli()
