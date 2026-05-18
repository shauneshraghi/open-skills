"""
benchmark.py — Performance benchmarks for docx creation and editing.

Measures time for:
- Creating documents of varying sizes (10, 100, 1000 paragraphs)
- Opening and saving an existing document
- Inserting and reading comments
- Inserting and accepting/rejecting tracked changes

Design: stdlib timeit for microsecond-resolution timing.
"""

from __future__ import annotations

import io
import json
import sys
import timeit
from pathlib import Path

from docx import Document
from docx.shared import Inches

sys.path.insert(0, str(Path(__file__).parent))
import create_docx
import edit_docx


def _make_png_bytes() -> bytes:
    """Return a minimal 1×1 white PNG for embedding in benchmarks."""
    import struct, zlib

    def chunk(tag: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", crc)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw  = b"\x00\xFF\xFF\xFF"
    idat = zlib.compress(raw)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", idat)
        + chunk(b"IEND", b"")
    )


PNG_BYTES = _make_png_bytes()


def bench_create(n_paragraphs: int, repeat: int = 3) -> float:
    """Return average seconds to create and save a document with n_paragraphs."""

    def run():
        doc = create_docx.create_document(title="Benchmark", author="bench")
        for i in range(n_paragraphs):
            create_docx.add_paragraph(doc, f"Paragraph {i}: The quick brown fox jumps over the lazy dog.")
        buf = io.BytesIO()
        doc.save(buf)

    times = timeit.repeat(run, number=1, repeat=repeat)
    return min(times)


def bench_open_save(path: str, repeat: int = 5) -> float:
    """Return average seconds to open and re-save an existing document."""

    def run():
        doc = edit_docx.open_document(path)
        buf = io.BytesIO()
        doc.save(buf)

    times = timeit.repeat(run, number=1, repeat=repeat)
    return min(times)


def bench_comments(n_comments: int, repeat: int = 3) -> float:
    """Return average seconds to insert n_comments into a document."""

    def run():
        doc = Document()
        for i in range(n_comments):
            para = doc.add_paragraph(f"Paragraph {i}")
            run_ = para.runs[0] if para.runs else para.add_run("text")
            doc.add_comment(run_, text=f"Comment {i}", author="bench")
        buf = io.BytesIO()
        doc.save(buf)

    times = timeit.repeat(run, number=1, repeat=repeat)
    return min(times)


def bench_track_changes(n_changes: int, repeat: int = 3) -> float:
    """Return average seconds to insert n tracked insertions and accept them."""

    def run():
        doc = Document()
        for i in range(n_changes):
            doc.add_paragraph(f"Base paragraph {i}")
        for i in range(n_changes):
            edit_docx.insert_tracked_insertion(doc, i, f" tracked {i}", author="bench")
        edit_docx.accept_all_changes(doc)
        buf = io.BytesIO()
        doc.save(buf)

    times = timeit.repeat(run, number=1, repeat=repeat)
    return min(times)


def bench_inline_images(n_images: int, repeat: int = 3) -> float:
    """Return average seconds to embed n inline images."""

    def run():
        doc = Document()
        for _ in range(n_images):
            para = doc.add_paragraph()
            run_ = para.add_run()
            run_.add_picture(io.BytesIO(PNG_BYTES), width=Inches(1))
        buf = io.BytesIO()
        doc.save(buf)

    times = timeit.repeat(run, number=1, repeat=repeat)
    return min(times)


def run_all(fixture_path: str | None = None) -> dict:
    """Run all benchmarks and return results as a dict."""
    results: dict = {}

    print("Benchmarking document creation...", file=sys.stderr)
    for n in (10, 100, 1000):
        key = f"create_{n}_paragraphs"
        t = bench_create(n)
        results[key] = round(t, 4)
        print(f"  {key}: {t:.4f}s", file=sys.stderr)

    if fixture_path:
        print("Benchmarking open/save...", file=sys.stderr)
        t = bench_open_save(fixture_path)
        results["open_save_fixture"] = round(t, 4)
        print(f"  open_save_fixture: {t:.4f}s", file=sys.stderr)

    print("Benchmarking comments...", file=sys.stderr)
    for n in (5, 50):
        key = f"comments_{n}"
        t = bench_comments(n)
        results[key] = round(t, 4)
        print(f"  {key}: {t:.4f}s", file=sys.stderr)

    print("Benchmarking track changes...", file=sys.stderr)
    for n in (5, 50):
        key = f"track_changes_{n}"
        t = bench_track_changes(n)
        results[key] = round(t, 4)
        print(f"  {key}: {t:.4f}s", file=sys.stderr)

    print("Benchmarking inline images...", file=sys.stderr)
    for n in (3, 10):
        key = f"inline_images_{n}"
        t = bench_inline_images(n)
        results[key] = round(t, 4)
        print(f"  {key}: {t:.4f}s", file=sys.stderr)

    return results


def _cli() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Run docx benchmarks")
    parser.add_argument("--fixture", help="Existing .docx for open/save benchmark")
    parser.add_argument("--json", action="store_true", help="Output JSON")
    args = parser.parse_args()

    results = run_all(fixture_path=args.fixture)

    if args.json:
        print(json.dumps(results, indent=2))
    else:
        print("\n=== Benchmark Results ===")
        for k, v in results.items():
            print(f"  {k:<35} {v:.4f}s")


if __name__ == "__main__":
    _cli()
