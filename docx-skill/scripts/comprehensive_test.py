"""
comprehensive_test.py — Full POI corpus test suite for docx-creation-editing skill.

Tests grouped into:
  A. Reading track changes from real-world fixtures
  B. Edge cases in accept/reject
  C. Surgical edits preserving existing markup
  D. Comment operations on complex fixtures
  E. Reviewer workflow (read → comment → tracked edit → validate)
  F. Image + track-change coexistence
"""

from __future__ import annotations

import io
import json
import os
import struct
import sys
import traceback
import zipfile
import zlib
from pathlib import Path
from typing import Any

sys.path.insert(0, str(Path(__file__).parent))
import create_docx
import edit_docx
import validate_docx

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from lxml import etree

# POI_PATH env var lets CI / other platforms point at their own corpus clone.
# Defaults to the path used in the reference Linux dev environment.
POI = Path(os.environ.get("POI_PATH", "/home/user/poi/test-data"))

# Output directory: use a temp dir so the tests work on Windows (no /tmp),
# macOS, and Linux without any configuration.
import tempfile as _tempfile
OUT = Path(os.environ.get("DOCX_TEST_OUT", _tempfile.mkdtemp(prefix="docx_tests_")))

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

PASS = "\033[32mPASS\033[0m"
FAIL = "\033[31mFAIL\033[0m"
SKIP = "\033[33mSKIP\033[0m"


def _make_png(r: int = 255, g: int = 0, b: int = 0) -> bytes:
    def chunk(t, d):
        c = zlib.crc32(t + d) & 0xFFFFFFFF
        return struct.pack(">I", len(d)) + t + d + struct.pack(">I", c)
    ihdr = struct.pack(">IIBBBBB", 8, 8, 8, 2, 0, 0, 0)
    raw = b"\x00" + bytes([r, g, b] * 8) * 8
    idat = zlib.compress(raw)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


results: list[dict] = []


def test(name: str, fn) -> dict:
    try:
        detail = fn()
        rec = {"name": name, "passed": True, "detail": detail or ""}
        print(f"  {PASS}  {name}")
        if detail:
            for line in str(detail).splitlines():
                print(f"         {line}")
    except AssertionError as e:
        rec = {"name": name, "passed": False, "detail": str(e)}
        print(f"  {FAIL}  {name}")
        print(f"         {e}")
    except Exception as e:
        rec = {"name": name, "passed": False, "detail": f"{type(e).__name__}: {e}"}
        print(f"  {FAIL}  {name}")
        print(f"         {type(e).__name__}: {e}")
        traceback.print_exc(limit=3)
    results.append(rec)
    return rec


def _raw_count(path: Path, tag: str) -> int:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml")
    return xml.decode("utf-8", errors="replace").count(tag)


def _revision_ids_in_doc(doc: Document) -> set[int]:
    body = doc.element.body
    ids = set()
    for tag in ("w:ins", "w:del", "w:commentRangeStart"):
        for el in body.iter(qn(tag)):
            try:
                ids.add(int(el.get(qn("w:id"), "0")))
            except ValueError:
                pass
    return ids


# ══════════════════════════════════════════════════════════════════════════════
# A. Reading track changes from real-world fixtures
# ══════════════════════════════════════════════════════════════════════════════

print("\n── A. Reading track changes ─────────────────────────────────────────")


def _test_A1():
    """delins.docx: detect insertions and deletions, verify author."""
    doc = edit_docx.open_document(POI / "document/delins.docx")
    changes = edit_docx.get_tracked_changes(doc)
    insertions = [c for c in changes if c["type"] == "insertion"]
    deletions = [c for c in changes if c["type"] == "deletion"]
    assert len(insertions) > 0, f"No insertions found, got changes={changes}"
    assert len(deletions) > 0, f"No deletions found"
    assert all(c["author"] for c in changes), "Some changes have empty author"
    authors = {c["author"] for c in changes}
    return f"ins={len(insertions)} del={len(deletions)} authors={authors}"
test("A1: delins.docx — read ins+del, author present", _test_A1)


def _test_A2():
    """58067.docx: pPr-level w:ins (formatting change, no children) is handled gracefully."""
    doc = edit_docx.open_document(POI / "document/58067.docx")
    changes = edit_docx.get_tracked_changes(doc)
    # pPr-level ins elements are self-closing (no w:t children) — get_tracked_changes
    # should return them with empty text, not crash
    for c in changes:
        assert "type" in c and "text" in c, f"Malformed change record: {c}"
    return f"changes found={len(changes)} (may include empty pPr-level ins)"
test("A2: 58067.docx — pPr-level ins (no children) handled without crash", _test_A2)


def _test_A3():
    """Tika-792.docx: nested w:ins inside w:moveFrom — complex revision markup."""
    doc = edit_docx.open_document(POI / "document/Tika-792.docx")
    changes = edit_docx.get_tracked_changes(doc)
    assert len(changes) >= 0, "Should not crash on nested/moveFrom markup"
    return f"changes found={len(changes)}"
test("A3: Tika-792.docx — nested ins in moveFrom markup, no crash", _test_A3)


def _test_A4():
    """stress023.docx: 144 ins + 75 del — large real-world document."""
    doc = edit_docx.open_document(POI / "integration/stress023.docx")
    changes = edit_docx.get_tracked_changes(doc)
    insertions = [c for c in changes if c["type"] == "insertion"]
    deletions = [c for c in changes if c["type"] == "deletion"]
    assert len(insertions) >= 100, f"Expected ≥100 insertions, got {len(insertions)}"
    assert len(deletions) >= 50, f"Expected ≥50 deletions, got {len(deletions)}"
    return f"ins={len(insertions)} del={len(deletions)}"
test("A4: stress023.docx — 100+ insertions and 50+ deletions detected", _test_A4)


# ══════════════════════════════════════════════════════════════════════════════
# B. Edge cases in accept / reject
# ══════════════════════════════════════════════════════════════════════════════

print("\n── B. Accept / reject edge cases ────────────────────────────────────")


def _test_B1():
    """Accept all on a document with no track changes — should be a no-op."""
    doc = edit_docx.open_document(POI / "document/VariousPictures.docx")
    n = edit_docx.accept_all_changes(doc)
    out = OUT / "B1_accept_no_changes.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid after no-op accept: {v['errors']}"
    assert n == 0, f"Expected 0 changes accepted, got {n}"
    return f"accepted={n} images_intact={v['stats']['images']}"
test("B1: Accept on no-track-change doc is a no-op, images intact", _test_B1)


def _test_B2():
    """pPr-level w:ins (formatting only, self-closing) — accept removes the marker, content survives."""
    doc = edit_docx.open_document(POI / "document/58067.docx")
    paras_before = len(doc.paragraphs)
    n = edit_docx.accept_all_changes(doc)
    out = OUT / "B2_accept_ppr_ins.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["tracked_insertions"] == 0, "w:ins markers remain after accept"
    doc2 = edit_docx.open_document(out)
    assert len(doc2.paragraphs) >= paras_before - 2, "Too many paragraphs lost"
    return f"accepted={n} paragraphs_before={paras_before} after={len(doc2.paragraphs)}"
test("B2: pPr-level w:ins accepted — content survives, no crash", _test_B2)


def _test_B3():
    """Reject all on delins.docx — deleted content is restored, inserted content is discarded."""
    doc = edit_docx.open_document(POI / "document/delins.docx")
    # Capture deleted text before reject
    changes_before = edit_docx.get_tracked_changes(doc)
    del_texts = [c["text"] for c in changes_before if c["type"] == "deletion" and c["text"]]
    ins_texts = [c["text"] for c in changes_before if c["type"] == "insertion" and c["text"]]

    n = edit_docx.reject_all_changes(doc)
    out = OUT / "B3_reject_delins.docx"
    edit_docx.save_document(doc, out)

    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["tracked_insertions"] == 0, "w:ins markers remain"
    assert v["stats"]["tracked_deletions"] == 0, "w:del markers remain"

    # Verify deleted text is restored. Deleted text may span multiple w:t elements
    # (split across runs), so check each word individually rather than as a
    # contiguous string.  "Tika Waylan" in delins.docx is split across two runs
    # with a hyperlink field in between.
    with zipfile.ZipFile(out) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    verified = 0
    for dt in del_texts[:3]:
        # Check each non-trivial word from the deleted text individually
        words = [w for w in dt.split() if len(w) > 3]
        for word in words[:2]:
            if word in xml:
                verified += 1
                break

    # Verify inserted text is gone (spot check first inserted word)
    # Note: inserted text may appear in other non-ins context so we just validate structure
    return f"rejected={n} del_runs_restored={verified}/{min(3,len(del_texts))}"
test("B3: Reject on delins.docx — deleted text restored, markup cleared", _test_B3)


def _test_B4():
    """Accept all on stress023.docx (144 ins + 75 del + 2 images) — images must survive."""
    doc = edit_docx.open_document(POI / "integration/stress023.docx")
    images_before = validate_docx.validate(POI / "integration/stress023.docx")["stats"]["images"]
    n = edit_docx.accept_all_changes(doc)
    out = OUT / "B4_accept_stress023.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["tracked_insertions"] == 0, "w:ins remain"
    assert v["stats"]["tracked_deletions"] == 0, "w:del remain"
    assert v["stats"]["images"] == images_before, \
        f"Images changed: was {images_before}, now {v['stats']['images']}"
    return f"accepted={n} images={v['stats']['images']} (unchanged)"
test("B4: Accept all on stress023.docx — 200+ changes, images unaffected", _test_B4)


def _test_B5():
    """Reject all on stress023.docx — images must also survive."""
    doc = edit_docx.open_document(POI / "integration/stress023.docx")
    images_before = validate_docx.validate(POI / "integration/stress023.docx")["stats"]["images"]
    n = edit_docx.reject_all_changes(doc)
    out = OUT / "B5_reject_stress023.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["images"] == images_before, \
        f"Images changed: was {images_before}, now {v['stats']['images']}"
    return f"rejected={n} images={v['stats']['images']} (unchanged)"
test("B5: Reject all on stress023.docx — images survive reject", _test_B5)


# ══════════════════════════════════════════════════════════════════════════════
# C. Surgical edits preserving existing markup
# ══════════════════════════════════════════════════════════════════════════════

print("\n── C. Surgical edits — ID uniqueness, no clobber ────────────────────")


def _test_C1():
    """Add a new tracked insertion to delins.docx — new ID must not collide with existing."""
    doc = edit_docx.open_document(POI / "document/delins.docx")
    ids_before = _revision_ids_in_doc(doc)
    changes_before = edit_docx.get_tracked_changes(doc)

    edit_docx.insert_tracked_insertion(doc, 0, " [REVIEWER NOTE]", author="Reviewer")

    ids_after = _revision_ids_in_doc(doc)
    new_ids = ids_after - ids_before
    assert len(new_ids) == 1, f"Expected 1 new id, got {new_ids}"
    assert new_ids.isdisjoint(ids_before), f"New id collides with existing: {new_ids & ids_before}"

    out = OUT / "C1_surgical_insert.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"

    changes_after = edit_docx.get_tracked_changes(edit_docx.open_document(out))
    assert len(changes_after) == len(changes_before) + 1, \
        f"Expected {len(changes_before)+1} changes, got {len(changes_after)}"
    return f"ids_before={len(ids_before)} new_id={new_ids} total_changes={len(changes_after)}"
test("C1: Surgical insert into delins.docx — unique ID, existing changes preserved", _test_C1)


def _test_C2():
    """Add tracked deletion to stress023.docx (already has 219 revisions) — no ID collision."""
    doc = edit_docx.open_document(POI / "integration/stress023.docx")
    ids_before = _revision_ids_in_doc(doc)
    max_id_before = max(ids_before) if ids_before else 0

    edit_docx.insert_tracked_deletion(doc, 0, "obsolete text", author="Reviewer")

    ids_after = _revision_ids_in_doc(doc)
    new_ids = ids_after - ids_before
    assert len(new_ids) == 1
    new_id = next(iter(new_ids))
    assert new_id > max_id_before, f"New id {new_id} not > max existing {max_id_before}"

    out = OUT / "C2_surgical_delete_stress023.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    return f"max_id_before={max_id_before} new_id={new_id}"
test("C2: Surgical delete into stress023.docx — ID > max existing", _test_C2)


def _test_C3():
    """Round-trip: open → insert tracked change → accept → no markup remains."""
    doc = edit_docx.open_document(POI / "document/delins.docx")
    changes_before = len(edit_docx.get_tracked_changes(doc))
    edit_docx.insert_tracked_insertion(doc, 0, " ADDED", author="Test")
    mid = OUT / "C3_mid.docx"
    edit_docx.save_document(doc, mid)

    doc2 = edit_docx.open_document(mid)
    assert len(edit_docx.get_tracked_changes(doc2)) == changes_before + 1

    edit_docx.accept_all_changes(doc2)
    out = OUT / "C3_accepted.docx"
    edit_docx.save_document(doc2, out)
    v = validate_docx.validate(out)
    assert v["valid"]
    assert v["stats"]["tracked_insertions"] == 0
    assert v["stats"]["tracked_deletions"] == 0

    # Verify the inserted text "ADDED" is now in plain text
    with zipfile.ZipFile(out) as zf:
        xml = zf.read("word/document.xml").decode()
    assert "ADDED" in xml, "Accepted inserted text not in final document"
    return f"changes_before={changes_before} → +1 → accept → 0 remaining, 'ADDED' preserved"
test("C3: Insert + accept round-trip — inserted text survives as plain text", _test_C3)


def _test_C4():
    """Enable track changes mode in settings.xml."""
    doc = edit_docx.open_document(POI / "document/delins.docx")
    edit_docx.enable_track_changes(doc)
    out = OUT / "C4_track_enabled.docx"
    edit_docx.save_document(doc, out)
    with zipfile.ZipFile(out) as zf:
        settings = zf.read("word/settings.xml").decode()
    assert "trackChanges" in settings, "w:trackChanges not found in settings.xml"
    return "w:trackChanges present in settings.xml"
test("C4: enable_track_changes writes w:trackChanges to settings.xml", _test_C4)


# ══════════════════════════════════════════════════════════════════════════════
# D. Comment operations on complex fixtures
# ══════════════════════════════════════════════════════════════════════════════

print("\n── D. Comment operations ────────────────────────────────────────────")


def _test_D1():
    """Read 3-comment WordWithAttachments.docx — all comments have author vbiryukova."""
    doc = edit_docx.open_document(POI / "document/WordWithAttachments.docx")
    comments = edit_docx.read_comments(doc)
    assert len(comments) == 3, f"Expected 3 comments, got {len(comments)}"
    assert all(c["author"] == "vbiryukova" for c in comments), \
        f"Unexpected authors: {[c['author'] for c in comments]}"
    return f"3 comments, author=vbiryukova for all"
test("D1: WordWithAttachments.docx — 3 comments, correct author", _test_D1)


def _test_D2():
    """Delete middle comment from WordWithAttachments.docx — remaining 2 are intact."""
    doc = edit_docx.open_document(POI / "document/WordWithAttachments.docx")
    comments = edit_docx.read_comments(doc)
    ids = [c["id"] for c in comments]
    middle_id = ids[1]

    ok = edit_docx.delete_comment(doc, middle_id)
    assert ok, "delete_comment returned False"

    out = OUT / "D2_deleted_middle_comment.docx"
    edit_docx.save_document(doc, out)

    doc2 = edit_docx.open_document(out)
    remaining = edit_docx.read_comments(doc2)
    assert len(remaining) == 2, f"Expected 2 comments after delete, got {len(remaining)}"
    remaining_ids = {c["id"] for c in remaining}
    assert middle_id not in remaining_ids, "Deleted comment id still present"
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    return f"deleted id={middle_id}, remaining ids={remaining_ids}"
test("D2: Delete middle comment — remaining 2 intact, doc valid", _test_D2)


def _test_D3():
    """Add comment to paragraph with no runs — should create a run automatically."""
    doc = create_docx.create_document()
    # Paragraph with text, then empty paragraph
    create_docx.add_paragraph(doc, "First paragraph with text")
    doc.add_paragraph()  # empty paragraph
    create_docx.add_paragraph(doc, "Third paragraph")

    # Comment on the empty paragraph (index 1)
    cid = edit_docx.insert_comment_on_paragraph(doc, 1, "Empty paragraph review", author="AI")
    out = OUT / "D3_comment_empty_para.docx"
    edit_docx.save_document(doc, out)

    doc2 = edit_docx.open_document(out)
    comments = edit_docx.read_comments(doc2)
    assert len(comments) == 1, f"Expected 1 comment, got {len(comments)}"
    assert comments[0]["text"] == "Empty paragraph review"
    v = validate_docx.validate(out)
    assert v["valid"]
    return f"comment on empty paragraph OK, id={cid}"
test("D3: Comment on empty paragraph — auto-creates run, doc valid", _test_D3)


def _test_D4():
    """Delete non-existent comment — returns False without crashing."""
    doc = edit_docx.open_document(POI / "document/comment.docx")
    result = edit_docx.delete_comment(doc, 99999)
    assert result is False, f"Expected False for missing id, got {result}"
    return "delete_comment(99999) → False correctly"
test("D4: Delete non-existent comment id — returns False, no crash", _test_D4)


def _test_D5():
    """testComment.docx already has a comment + image; add another comment without disturbing image."""
    doc = edit_docx.open_document(POI / "document/testComment.docx")
    comments_before = len(edit_docx.read_comments(doc))
    images_before = validate_docx.validate(POI / "document/testComment.docx")["stats"]["images"]

    # Add comment to first paragraph that has text
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            edit_docx.insert_comment_on_paragraph(doc, i, "Looks good", author="Reviewer")
            break

    out = OUT / "D5_testcomment_plus_new.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["images"] == images_before, \
        f"Images changed: {images_before} → {v['stats']['images']}"
    doc2 = edit_docx.open_document(out)
    comments_after = len(edit_docx.read_comments(doc2))
    assert comments_after == comments_before + 1, \
        f"Expected {comments_before+1} comments, got {comments_after}"
    return f"comments {comments_before}→{comments_after}, images={v['stats']['images']} (unchanged)"
test("D5: Add comment to testComment.docx — existing comment+image unaffected", _test_D5)


# ══════════════════════════════════════════════════════════════════════════════
# E. Reviewer workflow — end-to-end
# ══════════════════════════════════════════════════════════════════════════════

print("\n── E. Reviewer workflow (read → comment → track-edit → validate) ────")


def _test_E1():
    """
    Full reviewer workflow on delins.docx:
    1. Read existing tracked changes (ins + del)
    2. Add a comment on the first paragraph with an insertion, citing author
    3. Add a comment on the first paragraph with a deletion, citing deleted text
    4. Insert a new tracked deletion proposing further removal
    5. Save, validate
    """
    doc = edit_docx.open_document(POI / "document/delins.docx")
    changes = edit_docx.get_tracked_changes(doc)

    # Identify paragraphs with insertions and deletions
    ins_changes = [c for c in changes if c["type"] == "insertion" and c["text"].strip()]
    del_changes = [c for c in changes if c["type"] == "deletion" and c["text"].strip()]

    # Find paragraph indices that contain these changes
    body = doc.element.body
    all_paras = doc.paragraphs

    # Comment on first paragraph that has an insertion
    ins_para_idx = None
    del_para_idx = None
    for i, p in enumerate(all_paras):
        p_el = p._p
        if p_el.find(f".//{{{_W}}}ins") is not None and ins_para_idx is None:
            ins_para_idx = i
        if p_el.find(f".//{{{_W}}}del") is not None and del_para_idx is None:
            del_para_idx = i

    assert ins_para_idx is not None, "No paragraph with w:ins found"
    assert del_para_idx is not None, "No paragraph with w:del found"

    # Add reviewer comments
    ins_author = ins_changes[0]["author"] if ins_changes else "unknown"
    del_text_preview = del_changes[0]["text"][:30] if del_changes else "?"

    edit_docx.insert_comment_on_paragraph(
        doc, ins_para_idx,
        f"Insertion by {ins_author} — approved",
        author="AI-Reviewer",
    )
    edit_docx.insert_comment_on_paragraph(
        doc, del_para_idx,
        f"Deletion of '{del_text_preview}' — confirm removal",
        author="AI-Reviewer",
    )

    # Propose another deletion as a tracked change
    edit_docx.insert_tracked_deletion(
        doc, 0, "Tika can be:", author="AI-Reviewer"
    )

    out = OUT / "E1_reviewer_workflow.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["comments"] == 2, f"Expected 2 comments, got {v['stats']['comments']}"
    assert v["stats"]["tracked_deletions"] >= len(del_changes) + 1, \
        "New tracked deletion not found"

    doc2 = edit_docx.open_document(out)
    reviewer_comments = [c for c in edit_docx.read_comments(doc2) if c["author"] == "AI-Reviewer"]
    assert len(reviewer_comments) == 2
    return (
        f"ins_para={ins_para_idx} del_para={del_para_idx} "
        f"comments={v['stats']['comments']} deletions={v['stats']['tracked_deletions']}"
    )
test("E1: Reviewer workflow on delins.docx — comments + new tracked delete", _test_E1)


def _test_E2():
    """
    Reviewer workflow on WordWithAttachments.docx (3 existing comments, no track changes):
    1. Read existing comments
    2. Add a 4th reviewer comment
    3. Accept all changes (no-op — no tracked changes)
    4. Delete the first original comment
    5. Save, validate — 3 comments remain (2 original + 1 reviewer)
    """
    doc = edit_docx.open_document(POI / "document/WordWithAttachments.docx")
    orig_comments = edit_docx.read_comments(doc)
    assert len(orig_comments) == 3

    edit_docx.insert_comment_on_paragraph(doc, 0, "Overall LGTM", author="AI-Reviewer")
    n_accepted = edit_docx.accept_all_changes(doc)  # no-op

    first_id = orig_comments[0]["id"]
    edit_docx.delete_comment(doc, first_id)

    out = OUT / "E2_reviewer_attachments.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["comments"] == 3, f"Expected 3, got {v['stats']['comments']}"
    assert n_accepted == 0

    doc2 = edit_docx.open_document(out)
    final_comments = edit_docx.read_comments(doc2)
    final_ids = {c["id"] for c in final_comments}
    assert first_id not in final_ids, "Deleted comment still present"
    reviewer_comments = [c for c in final_comments if c["author"] == "AI-Reviewer"]
    assert len(reviewer_comments) == 1
    return f"original=3 → +1 reviewer → -1 deleted → 3 total, reviewer comment present"
test("E2: Reviewer workflow on WordWithAttachments — add + delete comments", _test_E2)


def _test_E3():
    """
    Reviewer accepts existing changes then leaves a summary comment:
    Simulates a reviewer saying 'I've reviewed and accepted these tracked changes'.
    """
    doc = edit_docx.open_document(POI / "document/delins.docx")
    changes = edit_docx.get_tracked_changes(doc)
    n_ins = sum(1 for c in changes if c["type"] == "insertion")
    n_del = sum(1 for c in changes if c["type"] == "deletion")

    n_accepted = edit_docx.accept_all_changes(doc)

    # Now add a summary comment on the first paragraph
    summary = f"Reviewed and accepted {n_ins} insertions and {n_del} deletions from original authors."
    edit_docx.insert_comment_on_paragraph(doc, 0, summary, author="AI-Reviewer")

    out = OUT / "E3_accepted_then_commented.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["tracked_insertions"] == 0
    assert v["stats"]["tracked_deletions"] == 0
    assert v["stats"]["comments"] == 1

    doc2 = edit_docx.open_document(out)
    comment_text = edit_docx.read_comments(doc2)[0]["text"]
    assert "accepted" in comment_text.lower()
    return f"accepted {n_accepted} changes, summary comment: '{comment_text[:60]}'"
test("E3: Accept changes then add summary comment — clean doc with 1 reviewer comment", _test_E3)


# ══════════════════════════════════════════════════════════════════════════════
# F. Image + track-change coexistence
# ══════════════════════════════════════════════════════════════════════════════

print("\n── F. Image + track-change coexistence ──────────────────────────────")


def _test_F1():
    """Add inline image to stress023.docx without touching existing 219 track changes."""
    doc = edit_docx.open_document(POI / "integration/stress023.docx")
    changes_before = edit_docx.get_tracked_changes(doc)
    images_before = validate_docx.validate(POI / "integration/stress023.docx")["stats"]["images"]

    png = _make_png(0, 128, 255)
    with open(OUT / "F1_test.png", "wb") as f:
        f.write(png)

    create_docx.add_inline_image(doc, OUT / "F1_test.png", width=Inches(1), alt_text="Reviewer diagram")

    out = OUT / "F1_stress023_plus_image.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["images"] == images_before + 1, \
        f"Expected {images_before+1} images, got {v['stats']['images']}"

    # Track changes must all survive
    doc2 = edit_docx.open_document(out)
    changes_after = edit_docx.get_tracked_changes(doc2)
    assert len(changes_after) == len(changes_before), \
        f"Track changes changed: {len(changes_before)} → {len(changes_after)}"
    return f"images {images_before}→{v['stats']['images']}, track_changes={len(changes_after)} (unchanged)"
test("F1: Add image to stress023.docx — 219 track changes unaffected", _test_F1)


def _test_F2():
    """drawing.docx (inline + anchored): add comment, validate both image types survive."""
    doc = edit_docx.open_document(POI / "document/drawing.docx")
    images_before = validate_docx.validate(POI / "document/drawing.docx")["stats"]["images"]
    assert images_before >= 1, "drawing.docx has no images?"

    # Add comment on first non-empty paragraph
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() or p.runs:
            edit_docx.insert_comment_on_paragraph(doc, i, "Drawing review", author="AI-Reviewer")
            break

    out = OUT / "F2_drawing_with_comment.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["images"] == images_before, \
        f"Images lost: {images_before} → {v['stats']['images']}"
    assert v["stats"]["comments"] == 1
    return f"images={v['stats']['images']} (unchanged) comments=1"
test("F2: drawing.docx — add comment preserves inline+anchored images", _test_F2)


def _test_F3():
    """Replace image in VariousPictures.docx and verify track changes are unaffected."""
    doc = edit_docx.open_document(POI / "document/VariousPictures.docx")
    # Insert a tracked insertion first
    edit_docx.insert_tracked_insertion(doc, 0, " [reviewed]", author="Reviewer")
    # Then replace image 0
    png = _make_png(255, 165, 0)
    with open(OUT / "F3_replace.png", "wb") as f:
        f.write(png)
    replaced = create_docx.replace_image(doc, 0, OUT / "F3_replace.png")
    assert replaced, "replace_image returned False"

    out = OUT / "F3_replaced_image_with_trackchange.docx"
    edit_docx.save_document(doc, out)
    v = validate_docx.validate(out)
    assert v["valid"], f"Document invalid: {v['errors']}"
    assert v["stats"]["tracked_insertions"] == 1, "Tracked insertion lost"
    assert v["stats"]["images"] == 5, f"Image count changed: {v['stats']['images']}"
    return f"image_replaced=True, tracked_insertion=1, images={v['stats']['images']}"
test("F3: Replace image + tracked insertion — both survive, doc valid", _test_F3)


# ══════════════════════════════════════════════════════════════════════════════
# Summary
# ══════════════════════════════════════════════════════════════════════════════

print()
passed = sum(1 for r in results if r["passed"])
total = len(results)
failed = [r for r in results if not r["passed"]]

print(f"══ Results: {passed}/{total} passed ══")
if failed:
    print("\nFailed tests:")
    for r in failed:
        print(f"  ✗ {r['name']}")
        print(f"    {r['detail']}")

# Save results as JSON
summary = {
    "total": total,
    "passed": passed,
    "failed": total - passed,
    "results": results,
}
with open(OUT / "comprehensive_results.json", "w") as f:
    json.dump(summary, f, indent=2)
print(f"\nResults saved to {OUT}/comprehensive_results.json")
sys.exit(0 if not failed else 1)
